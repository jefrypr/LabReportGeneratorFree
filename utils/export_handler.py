# utils/export_handler.py
"""
Handler untuk export laporan ke format DOCX dan HTML.
"""

import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Helpers ────────────────────────────────────────────────────────────────

def _set_font(run, name="Times New Roman", size=12):
    run.font.name = name
    run.font.size = Pt(size)


def _add_heading_styled(doc: Document, text: str, level: int = 1):
    """Tambahkan heading dengan styling manual agar lebih konsisten."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    _set_font(run, size=13 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_body_paragraph(doc: Document, text: str, justify: bool = True):
    """Tambahkan paragraf body dengan formatting standar laporan."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Handle italic markers (_kata_ → italic)
    parts = re.split(r'(_[^_]+_)', text)
    for part in parts:
        if part.startswith('_') and part.endswith('_') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        _set_font(run)

    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    return p


def _parse_sections(results: dict) -> dict:
    """Ambil section dari results dict."""
    return {
        "analysis": results.get("analysis", "").strip(),
        "conclusion": results.get("conclusion", "").strip(),
        "langkah_kerja": results.get("langkah_kerja", "").strip(),
    }


def _format_refs(refs_raw: str) -> list[str]:
    """Format daftar pustaka menjadi list dengan nomor."""
    lines = [r.strip() for r in refs_raw.strip().split("\n") if r.strip()]
    result = []
    for i, line in enumerate(lines, 1):
        cleaned = re.sub(r'^\[?\d+\]?\.?\s*', '', line)
        result.append(f"[{i}] {cleaned}")
    return result


# ─── DOCX Export ────────────────────────────────────────────────────────────

def export_to_docx(data: dict, results: dict) -> bytes:
    """
    Export seluruh hasil ke file DOCX berformat laporan praktikum.

    Args:
        data: dict berisi mata_praktikum, judul, daftar_pustaka
        results: dict berisi analysis, conclusion, langkah_kerja

    Returns:
        bytes: isi file .docx
    """
    doc = Document()
    sections = _parse_sections(results)

    # --- Margin setup ---
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)

    # --- Cover / Judul ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("LAPORAN PRAKTIKUM")
    title_run.bold = True
    _set_font(title_run, size=14)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"{data.get('mata_praktikum', '')} — {data.get('judul', '')}")
    sub_run.bold = True
    _set_font(sub_run, size=13)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Tanggal: {datetime.now().strftime('%d %B %Y')}")
    _set_font(date_run, size=11)

    doc.add_paragraph()  # Spacer

    # --- Langkah Kerja (jika ada) ---
    if sections["langkah_kerja"]:
        _add_heading_styled(doc, "LANGKAH KERJA")
        lines = [l.strip() for l in sections["langkah_kerja"].split("\n") if l.strip()]
        for line in lines:
            # Hapus prefix nomor yang mungkin ada
            clean = re.sub(r'^\d+[\.\)]\s*', '', line)
            if clean:
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(clean)
                _set_font(run)
                p.paragraph_format.space_after = Pt(4)

    # --- Analisis dan Pembahasan ---
    if sections["analysis"]:
        _add_heading_styled(doc, "ANALISIS DAN PEMBAHASAN")
        # Pisah per paragraf (double newline)
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', sections["analysis"]) if p.strip()]
        if not paragraphs:
            # Fallback: pisah per baris single
            paragraphs = [l.strip() for l in sections["analysis"].split("\n") if l.strip()]
        for para_text in paragraphs:
            _add_body_paragraph(doc, para_text)

    # --- Kesimpulan ---
    if sections["conclusion"]:
        _add_heading_styled(doc, "KESIMPULAN")
        lines = [l.strip() for l in sections["conclusion"].split("\n") if l.strip()]
        for line in lines:
            clean = re.sub(r'^\d+[\.\)]\s*', '', line)
            if clean:
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(clean)
                _set_font(run)
                p.paragraph_format.space_after = Pt(4)

    # --- Daftar Pustaka ---
    refs_raw = data.get("daftar_pustaka", "").strip()
    if refs_raw:
        _add_heading_styled(doc, "DAFTAR PUSTAKA")
        refs = _format_refs(refs_raw)
        for ref in refs:
            p = doc.add_paragraph()
            run = p.add_run(ref)
            _set_font(run, size=11)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
            p.paragraph_format.space_after = Pt(4)

    # --- Simpan ke bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── HTML Export ─────────────────────────────────────────────────────────────

def export_to_html(data: dict, results: dict) -> str:
    """
    Export seluruh hasil ke HTML yang dapat dibuka di browser.

    Args:
        data: dict berisi mata_praktikum, judul, daftar_pustaka
        results: dict berisi analysis, conclusion, langkah_kerja

    Returns:
        str: konten HTML lengkap
    """
    sections = _parse_sections(results)

    def md_to_html_inline(text: str) -> str:
        """Convert _italic_ markers ke <em> tag."""
        return re.sub(r'_([^_]+)_', r'<em>\1</em>', text)

    # --- Analysis HTML ---
    analysis_html = ""
    if sections["analysis"]:
        paras = [p.strip() for p in re.split(r'\n\s*\n', sections["analysis"]) if p.strip()]
        if not paras:
            paras = [l.strip() for l in sections["analysis"].split("\n") if l.strip()]
        analysis_html = "\n".join(
            f"<p>{md_to_html_inline(p)}</p>" for p in paras
        )

    # --- Conclusion HTML ---
    conc_html = ""
    if sections["conclusion"]:
        lines = [l.strip() for l in sections["conclusion"].split("\n") if l.strip()]
        items = []
        for line in lines:
            clean = re.sub(r'^\d+[\.\)]\s*', '', line)
            if clean:
                items.append(f"<li>{md_to_html_inline(clean)}</li>")
        conc_html = "<ol>\n" + "\n".join(items) + "\n</ol>" if items else ""

    # --- Langkah Kerja HTML ---
    lk_html = ""
    if sections["langkah_kerja"]:
        lines = [l.strip() for l in sections["langkah_kerja"].split("\n") if l.strip()]
        items = []
        for line in lines:
            clean = re.sub(r'^\d+[\.\)]\s*', '', line)
            if clean:
                items.append(f"<li>{md_to_html_inline(clean)}</li>")
        lk_html = "<ol>\n" + "\n".join(items) + "\n</ol>" if items else ""

    # --- References HTML ---
    refs_html = ""
    refs_raw = data.get("daftar_pustaka", "").strip()
    if refs_raw:
        refs = _format_refs(refs_raw)
        refs_html = "\n".join(f"<p class='ref'>{r}</p>" for r in refs)

    # --- Sections assembly ---
    lk_section = f"""
    <div class="section">
        <h2>Langkah Kerja</h2>
        {lk_html}
    </div>""" if lk_html else ""

    refs_section = f"""
    <div class="section">
        <h2>Daftar Pustaka</h2>
        {refs_html}
    </div>""" if refs_html else ""

    html = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Laporan Praktikum — {data.get('mata_praktikum', '')}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.8;
            color: #1a1a1a;
            background: #fff;
            max-width: 210mm;
            margin: 0 auto;
            padding: 30mm 25mm 30mm 30mm;
        }}
        .cover {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid #333;
        }}
        .cover h1 {{
            font-size: 16pt;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        .cover h2 {{
            font-size: 13pt;
            font-weight: bold;
            margin-top: 8px;
        }}
        .cover .date {{
            font-size: 11pt;
            color: #555;
            margin-top: 8px;
        }}
        .section {{
            margin-top: 28px;
        }}
        h2 {{
            font-size: 13pt;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            border-bottom: 1px solid #ccc;
            padding-bottom: 4px;
            margin-bottom: 14px;
        }}
        p {{
            text-align: justify;
            text-indent: 1.25cm;
            margin-bottom: 8px;
        }}
        ol {{
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 6px;
            text-align: justify;
        }}
        .ref {{
            text-indent: -1.25cm;
            padding-left: 1.25cm;
            font-size: 10.5pt;
            margin-bottom: 4px;
        }}
        em {{ font-style: italic; }}
        @media print {{
            body {{ padding: 30mm 25mm 30mm 30mm; }}
        }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>Laporan Praktikum</h1>
        <h2>{data.get('mata_praktikum', '')} &mdash; {data.get('judul', '')}</h2>
        <div class="date">Tanggal: {datetime.now().strftime('%d %B %Y')}</div>
    </div>

    {lk_section}

    <div class="section">
        <h2>Analisis dan Pembahasan</h2>
        {analysis_html}
    </div>

    <div class="section">
        <h2>Kesimpulan</h2>
        {conc_html}
    </div>

    {refs_section}
</body>
</html>"""

    return html
